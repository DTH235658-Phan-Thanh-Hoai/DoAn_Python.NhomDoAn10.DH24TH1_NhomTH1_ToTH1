import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
import pyodbc
from docx import Document 
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

import frmTongQuan

# === TAB HÓA ĐƠN ===
class tabHoaDon(tk.Frame):
    def __init__(self, parent, conn, user, controller=None):
        super().__init__(parent, bg="white")

        # === CHUỖI KẾT NỐI ===
        self.conn = conn

        # === LƯU MÃ NHÂN VIÊN ===
        self.user = user

        # === LƯU THAM CHIẾU CONTROLLER ===
        self.controller = controller

        # === KHUNG TÌM KIẾM ===
        frame_search = tk.Frame(self, bg="#E3F2FD", padx=10, pady=10)
        frame_search.pack(fill="x", padx=20, pady=5)

        tk.Label(frame_search, text="Tìm kiếm:", font=("Segoe UI", 10), bg="#E3F2FD").pack(side="left", padx=5)
        self.txt_timkiem = tk.Entry(frame_search, font=("Segoe UI", 10), width=25)
        self.txt_timkiem.pack(side="left", padx=5)

        self.search_option = tk.StringVar(value="mahd")
        tk.Radiobutton(frame_search, text="Theo mã hóa đơn", variable=self.search_option, value="mahd", bg="#E3F2FD", font=("Segoe UI", 10)).pack(side="left", padx=10)
        tk.Radiobutton(frame_search, text="Theo mã khách hàng", variable=self.search_option, value="makh", bg="#E3F2FD", font=("Segoe UI", 10)).pack(side="left" , padx=10)
        tk.Radiobutton(frame_search, text="Theo trạng thái phiếu", variable=self.search_option, value="trangthai", bg="#E3F2FD", font=("Segoe UI", 10)).pack(side="left", padx=10)
        tk.Button(frame_search, text="Tìm", font=("Segoe UI", 10, "bold"), command=self.TimKiem, bg="#1565C0", fg="white", bd=0, padx=10, pady=5).pack(side="left", padx=10)
        tk.Button(frame_search, text="Hủy", font=("Segoe UI", 10, "bold"), command=self.HuyTimKiem, bg="#E53935", fg="white", bd=0, padx=10, pady=5).pack(side="left", padx=10)

        # === BẢNG DANH SÁCH HÓA ĐƠN ===
        frame_table = tk.Frame(self, bg="white")
        frame_table.pack(fill="both", expand=True, padx=20, pady=10)

        columns = ("MaHD", "NgayBan", "MaNV", "MaKH", "ThanhTien", "TrangThai")
        # --- Tạo Scrollbar ---
        scroll_y = ttk.Scrollbar(frame_table, orient="vertical")
        scroll_x = ttk.Scrollbar(frame_table, orient="horizontal")

        self.trHienThi = ttk.Treeview( frame_table, show="headings",  columns=columns, height=12, yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)

        # --- Gắn Scrollbar ---
        scroll_y.config(command=self.trHienThi.yview)
        scroll_x.config(command=self.trHienThi.xview)

        # --- Bố trí Scrollbar ---
        scroll_y.pack(side="right", fill="y")
        scroll_x.pack(side="bottom", fill="x")
        self.trHienThi.pack(fill="both", expand=True)

        self.trHienThi.heading("MaHD", text="Mã hóa đơn")
        self.trHienThi.heading("NgayBan", text="Ngày bán")
        self.trHienThi.heading("MaNV", text="Mã nhân viên")
        self.trHienThi.heading("MaKH", text="Mã khách hàng")
        self.trHienThi.heading("ThanhTien", text="Thành tiền")
        self.trHienThi.heading("TrangThai", text="Trạng thái")

        self.trHienThi.column("MaHD", width=100)
        self.trHienThi.column("NgayBan", width=100, anchor="center")
        self.trHienThi.column("MaNV", width=100)
        self.trHienThi.column("MaKH", width=100)
        self.trHienThi.column("ThanhTien", width=100, anchor="center")
        self.trHienThi.column("TrangThai", width=100)

        # Thêm style cho Treeview
        style = ttk.Style()
        style.configure("Treeview.Heading", font=("Segoe UI", 11, "bold"))
        style.configure("Treeview", font=("Segoe UI", 10), rowheight=28)

        # ==== Nút thao tác ====
        frame_btn = tk.Frame(self, bg="white")
        frame_btn.pack(pady=10)

        tk.Button(frame_btn, text="🧐 Xem chi tiết", bg="#EC9428", fg="white", font=("Segoe UI", 10, "bold"),command=self.XemChiTiet, padx=15, pady=5, bd=0).pack(side="left", padx=5)
        tk.Button(frame_btn, text="💳 Thanh toán hóa đơn", bg="#43A047", fg="white",  font=("Segoe UI", 10, "bold"), command=self.ThanhToanHoaDonBan, padx=15, pady=5, bd=0).pack(side="left", padx=5)
        tk.Button(frame_btn, text="❌ Hủy hóa đơn", bg="#E53935", fg="white",  font=("Segoe UI", 10, "bold"), padx=15, command=self.HuyHoaDonBan, pady=5, bd=0).pack(side="left", padx=5)
        tk.Button(frame_btn, text="🔄 Làm mới", bg="#1E88E5", fg="white", font=("Segoe UI", 10, "bold"), padx=15, command=self.load_hoa_don, pady=5, bd=0).pack(side="left", padx=5)
        tk.Button(frame_btn, text="🗑️ Xóa", bg="#B71C1C", fg="white", font=("Segoe UI", 10, "bold"), padx=15, command=self.XoaHoaDonVinhVien, pady=5, bd=0).pack(side="left", padx=5)
        tk.Button(frame_btn, text="🖨️ In", bg="#E51E9C", fg="white", font=("Segoe UI", 10, "bold"), padx=15, command=self.InHoaDon, pady=5, bd=0).pack(side="left", padx=5)

        # === TẢI DỮ LIỆU HÓA ĐƠN ===
        self.load_hoa_don()

    def load_hoa_don(self):
        try:
            self.trHienThi.delete(*self.trHienThi.get_children())
            cursor = self.conn.cursor()
            # Khởi tạo câu truy vấn cơ bản
            sql_query = "SELECT MaHD, NgayBan, MaNV, MaKH, TongTien, TrangThai FROM HOADONBAN"
            params = []
            
            # KIỂM TRA QUYỀN TRUY CẬP
            if self.user.lower() != "admin":
                # Nếu không phải admin, chỉ lấy hóa đơn của nhân viên đó
                sql_query = sql_query + " WHERE MaNV = ?"
                params.append(self.user)
            
            # Thực thi truy vấn
            cursor.execute(sql_query, params)

            for row in cursor.fetchall():
                ngay_ban = datetime.strptime(str(row.NgayBan).split(" ")[0], "%Y-%m-%d")

                formatted_row = (
                    row.MaHD, ngay_ban.strftime("%d/%m/%Y"), row.MaNV, row.MaKH,
                    f"{float(row.TongTien):,.0f}" if row.TongTien else "0", row.TrangThai
                )
                self.trHienThi.insert("", tk.END, values=formatted_row)

        except Exception as e:
            messagebox.showerror("Lỗi", "Không thể tải dữ liệu hóa đơn:\n" + str(e))

    # === Hàm xem chi tiết phiếu nhập ===
    def XemChiTiet(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 hóa đơn để xem chi tiết.")
            return

        ma_hd = self.trHienThi.item(selected[0], "values")[0]

        try:
            chitiethoadon= tk.Toplevel(self)
            chitiethoadon.title("Chi tiết hóa đơn: " + ma_hd)
            chitiethoadon.geometry("800x800")
            chitiethoadon.resizable(False, False)
            chitiethoadon.configure(bg="white")
            

            tk.Label(chitiethoadon, text="Chi tiết hóa đơn bán " + ma_hd, font=("Segoe UI", 12, "bold"), bg="white", fg="#0D47A1").pack(pady=10)
            
            cursor = self.conn.cursor()
            cursor.execute("""SELECT nv.MaNV, nv.TenNV, kh.MaKH, kh.TenKH, hdb.NgayBan, hdb.TongTien, hdb.TrangThai
                              FROM HOADONBAN hdb JOIN NHANVIEN nv ON hdb.MaNV = nv.MaNV
                              JOIN KHACHHANG kh ON hdb.MaKH = kh.MaKH
                              WHERE hdb.MaHD = ?""", (ma_hd,))
            thong_tin = cursor.fetchone()

            Frame_thongtin = tk.Frame(chitiethoadon, bg="white")
            Frame_thongtin.pack(pady=5)
            tk.Label(Frame_thongtin, text="Mã nhân viên:", font=("Segoe UI", 10), bg="white").grid(row=0, column=0, padx=10, sticky="w")
            tk.Label(Frame_thongtin, text=thong_tin.MaNV, font=("Segoe UI", 10, "bold"), bg="white", fg="#1565C0").grid(row=0, column=1, padx=10, sticky="w")

            tk.Label(Frame_thongtin, text="Tên nhân viên:", font=("Segoe UI", 10), bg="white").grid(row=0, column=2, padx=10, sticky="w")
            tk.Label(Frame_thongtin, text=thong_tin.TenNV, font=("Segoe UI", 10, "bold"), bg="white", fg="#1565C0").grid(row=0, column=3, padx=10, sticky="w")

            tk.Label(Frame_thongtin, text="Mã khách hàng:", font=("Segoe UI", 10), bg="white").grid(row=1, column=0, padx=10, sticky="w")
            tk.Label(Frame_thongtin, text=thong_tin.MaKH, font=("Segoe UI", 10, "bold"), bg="white", fg="#1565C0").grid(row=1, column=1, padx=10, sticky="w")

            tk.Label(Frame_thongtin, text="Tên khách hàng:", font=("Segoe UI", 10), bg="white").grid(row=1, column=2, padx=10, sticky="w")
            tk.Label(Frame_thongtin, text=thong_tin.TenKH, font=("Segoe UI", 10, "bold"), bg="white", fg="#1565C0").grid(row=1, column=3, padx=10, sticky="w")

            tk.Label(Frame_thongtin, text="Ngày bán:", font=("Segoe UI", 10), bg="white").grid(row=2, column=0, padx=10, sticky="w")
            
            ngay_nhap = datetime.strptime(str(thong_tin.NgayBan).split(" ")[0], "%Y-%m-%d")
            tk.Label(Frame_thongtin, text=ngay_nhap.strftime("%d/%m/%Y"), font=("Segoe UI", 10, "bold"), bg="white", fg="#43A047").grid(row=2, column=1, padx=10, sticky="w")

            tk.Label(Frame_thongtin, text="Trạng thái:", font=("Segoe UI", 10), bg="white").grid(row=2, column=2, padx=10, sticky="w")
            tk.Label(Frame_thongtin, text=thong_tin.TrangThai, font=("Segoe UI", 10, "bold"), bg="white", fg="#43A047").grid(row=2, column=3, padx=10, sticky="w")
            
            columns = ("MaCTHD", "MaTivi", "TenTivi", "SoLuong", "DonGia", "ThanhTien")
            tree = ttk.Treeview(chitiethoadon, columns=columns, show="headings", height=10)
            tree.pack(fill="both", expand=True, padx=15, pady=10)

            tree.heading("MaCTHD", text="Mã CTHD")
            tree.heading("MaTivi", text="Mã Tivi")
            tree.heading("TenTivi", text="Tên Tivi")
            tree.heading("SoLuong", text="Số Lượng")
            tree.heading("DonGia", text="Đơn Giá")
            tree.heading("ThanhTien", text="Thành Tiền")

            tree.column("MaCTHD", width=100)
            tree.column("MaTivi", width=100)
            tree.column("TenTivi", width=200)
            tree.column("SoLuong", width=100, anchor="center")
            tree.column("DonGia", width=100, anchor="center")
            tree.column("ThanhTien", width=100, anchor="center")

            tk.Label(chitiethoadon, text="Tổng tiền:", font=("Segoe UI", 10, "bold"), bg="white").pack(side="left", padx=20)
            tk.Label(chitiethoadon, text=f"{float(thong_tin.TongTien):,.0f} đ", font=("Segoe UI", 10, "bold"), bg="white", fg="red").pack(side="right", padx=20)
            
            cursor.execute("""
                SELECT cthd.MaCTHD, cthd.MaTivi, tv.TenTivi, cthd.SoLuong, cthd.DonGia, (cthd.SoLuong * cthd.DonGia) AS ThanhTien
                FROM CHITIETHOADON cthd
                JOIN TIVI tv ON cthd.MaTivi = tv.MaTivi
                WHERE cthd.MaHD = ?
            """, (ma_hd,))
            rows = cursor.fetchall()

            if not rows:
                messagebox.showinfo("Thông báo", "Không có chi tiết cho phiếu" + ma_hd)
                chitiethoadon.destroy()
                return

            for r in rows:
                tree.insert("", tk.END, values=(r.MaCTHD,r.MaTivi, r.TenTivi,r.SoLuong,f"{float(r.DonGia):,.0f} đ",f"{float(r.ThanhTien):,.0f} đ"))

            cursor.close()
        except Exception as e:
            messagebox.showerror("Lỗi", "Không thể xem chi tiết hóa đơn:\n" + str(e))

    def ThanhToanHoaDonBan(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 hoa đơn để duyệt.")
            return
        ma_hd= self.trHienThi.item(selected[0], "values")[0]
        # Bổ sung lấy MaNV và kiểm tra quyền nếu cần thiết ở đây

        trang_thai = self.trHienThi.item(selected[0], "values")[5]

        if trang_thai != "Chờ thanh toán":
            messagebox.showwarning("Thông báo", "Chỉ có thể thanh toán hóa đơn chờ thanh toán.")
            return
        
        traloi = messagebox.askyesno("Xác nhận", "Bạn có chắc chắn thanh toán hóa đơn " + ma_hd + " không?")
        if traloi:
            try:
                cursor = self.conn.cursor()

                # Lấy chi tiết hóa đơn VÀ TỒN KHO hiện tại
                cursor.execute("""
                    SELECT cthd.MaTivi, cthd.SoLuong, tv.SoLuongTon
                    FROM CHITIETHOADON cthd
                    JOIN TIVI tv ON cthd.MaTivi = tv.MaTivi
                    WHERE cthd.MaHD = ?
                """, (ma_hd,))
                chitiet = cursor.fetchall()
                
                # KIỂM TRA TỒN KHO TRƯỚC KHI THỰC HIỆN GIAO DỊCH 
                for item in chitiet:
                    ma_tivi = item.MaTivi
                    so_luong_ban = item.SoLuong
                    ton_kho_hien_tai = item.SoLuongTon
                    
                    if so_luong_ban > ton_kho_hien_tai:
                        messagebox.showerror("Lỗi tồn kho", f"Không đủ tồn kho để thanh toán! Tivi {ma_tivi} cần {so_luong_ban} nhưng kho chỉ còn {ton_kho_hien_tai}.")
                        cursor.close()
                        return # Dừng toàn bộ giao dịch ngay lập tức

                # THỰC HIỆN GIAO DỊCH (Nếu kiểm tra thành công)
                
                # Trừ kho trước
                for item in chitiet:
                    ma_tivi = item.MaTivi
                    so_luong = item.SoLuong
                    
                    cursor.execute("""
                        UPDATE TIVI
                        SET SoLuongTon = SoLuongTon - ?
                        WHERE MaTivi = ? """, (so_luong, ma_tivi))
                
                # Cập nhật trạng thái sau khi trừ kho thành công
                cursor.execute("""
                    UPDATE HOADONBAN
                    SET TrangThai = N'Đã thanh toán'
                    WHERE MaHD = ?
                """, (ma_hd,))
                
                self.conn.commit()
                cursor.close()
                messagebox.showinfo("Thành công", "Hóa đơn đã thanh toán thành công!")

                # Làm mới lại danh sách hóa đơn
                self.load_hoa_don()
                                
            except Exception as e:
                self.conn.rollback() # Hoàn tác mọi thay đổi nếu có lỗi
                messagebox.showerror("Lỗi", "Đã xảy ra lỗi khi thanh toán hóa đơn:\n" + str(e))

    def HuyHoaDonBan(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 hóa đơn để hủy")
            return
        
        ma_hd = self.trHienThi.item(selected[0], "values")[0]
        trang_thai = self.trHienThi.item(selected[0], "values")[5]

        if trang_thai != "Chờ thanh toán":
            messagebox.showwarning("Thông báo", "Chỉ có thể hủy hóa đơn chưa thanh toán.")
            return
        
        traloi = messagebox.askyesno("Xác nhận", "Bạn có chắc chắn muốn hủy hóa đơn " + ma_hd + " không?")
        if traloi:
            try:
                cursor = self.conn.cursor()
                cursor.execute("""
                    UPDATE HOADONBAN
                    SET TrangThai = N'Đã hủy'
                    WHERE MaHD = ?""", (ma_hd,))
                
                self.conn.commit()
                cursor.close()
                messagebox.showinfo("Thành công", "Hóa đơn đã được hủy thành công!")
                self.load_hoa_don()

            except Exception as e:
                messagebox.showerror("Lỗi", "Đã xảy ra lỗi khi hủy hóa đơn:\n" + str(e))


    def TimKiem(self):
        if self.txt_timkiem.get().strip() == "":
            messagebox.showwarning("Thông báo", "Vui lòng nhập từ khóa tìm kiếm.")
            return
        
        timkkiem = self.search_option.get().strip()
        self.trHienThi.delete(*self.trHienThi.get_children())

        try:
            cursor = self.conn.cursor()

            if timkkiem == "mahd":
                keyword = self.txt_timkiem.get()
                cursor.execute("""
                    SELECT MaHD, NgayBan, MaNV, MaKH, ThanhTien, TrangThai
                    FROM HOADONBAN
                    WHERE MaHD = ?""", (keyword,))
            
            elif timkkiem == "makh":
                keyword = self.txt_timkiem.get()
                cursor.execute("""
                    SELECT MaHD, NgayBan, MaNV, MaKH, ThanhTien, TrangThai
                    FROM HOADONBAN
                    WHERE MaKH = ?""", (keyword,))
            
            elif timkkiem == "trangthai":
                keyword = self.txt_timkiem.get()
                cursor.execute("""
                    SELECT MaHD, NgayBan, MaNV, MaKH, ThanhTien, TrangThai
                    FROM HOADONBAN
                    WHERE TrangThai = ?""", (keyword,))
                
            rows = cursor.fetchall()
            for row in rows:
                ngay_ban = datetime.strptime(str(row.NgayBan).split(" ")[0], "%Y-%m-%d")

                formatted_row = (
                    row.MaHD, ngay_ban.strftime("%d/%m/%y"), row.MaNV, row.MaKH,
                    f"{float(row.ThanhTien):,.0f}" if row.ThanhTien else "0", row.TrangThai
                )
                self.trHienThi.insert("", tk.END, values=formatted_row)
            cursor.close()

        except Exception as e:
            messagebox.showerror("Lỗi", "Không thể tìm kiếm hóa đơn:\n" + str(e))

    def HuyTimKiem(self):
        self.load_hoa_don()
        self.txt_timkiem.delete(0, tk.END)

    # === HÀM IN HÓA ĐƠN RA WORD ===
    def InHoaDon(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 hóa đơn để in.")
            return

        ma_hd = self.trHienThi.item(selected[0], "values")[0]
        
        # --- Lấy dữ liệu chi tiết ---
        try:
            cursor = self.conn.cursor()
            
            # 1. Lấy thông tin chung của hóa đơn
            cursor.execute("""SELECT hdb.MaHD, nv.MaNV, nv.TenNV, kh.MaKH, kh.TenKH, hdb.NgayBan, hdb.ThanhTien, hdb.TrangThai
                              FROM HOADONBAN hdb 
                              JOIN NHANVIEN nv ON hdb.MaNV = nv.MaNV
                              JOIN KHACHHANG kh ON hdb.MaKH = kh.MaKH
                              WHERE hdb.MaHD = ?""", (ma_hd,))
            thong_tin = cursor.fetchone()

            if not thong_tin:
                messagebox.showerror("Lỗi", f"Không tìm thấy thông tin hóa đơn {ma_hd}")
                return

            # 2. Lấy chi tiết các mặt hàng
            cursor.execute("""
                SELECT cthd.MaCTHD, cthd.MaTivi, tv.TenTivi, cthd.SoLuong, cthd.DonGia, (cthd.SoLuong * cthd.DonGia) AS ThanhTien
                FROM CHITIETHOADON cthd
                JOIN TIVI tv ON cthd.MaTivi = tv.MaTivi
                WHERE cthd.MaHD = ?
            """, (ma_hd,))
            chi_tiet_hang = cursor.fetchall()
            cursor.close()

            if not chi_tiet_hang:
                messagebox.showinfo("Thông báo", f"Hóa đơn {ma_hd} không có chi tiết hàng hóa.")
                return

        except Exception as e:
            messagebox.showerror("Lỗi CSDL", "Không thể lấy dữ liệu để in hóa đơn:\n" + str(e))
            return
            
        # --- Tạo File Word ---
        try:
            document = Document()
            
            # Đặt font và size mặc định cho document
            style = document.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(11)

            # Dữ liệu định dạng
            ngay_ban_date = datetime.strptime(str(thong_tin.NgayBan).split(" ")[0], "%Y-%m-%d")
            tong_tien_str = f"{float(thong_tin.ThanhTien):,.0f} đ"
            

            # Tiêu đề chính (Căn giữa)
            heading = document.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.add_run(f"HÓA ĐƠN BÁN HÀNG SỐ: {ma_hd}")
            run.bold = True
            run.font.size = Pt(16)
            document.add_paragraph() # Dòng trống
            
            # --- Bố trí Thông tin chung (Dùng Tab Stop cố định) ---
            
            # Thiết lập Tab Stop: Căn lề trái tại vị trí 8.0 cm
            TAB_STOP_POS = Cm(8.0) 
            
            # Thông tin Nhân viên
            p1 = document.add_paragraph()
            p1.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POS, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p1.add_run("Mã nhân viên: ").bold = False
            run = p1.add_run(thong_tin.MaNV)
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) 
            p1.add_run('\t') 
            p1.add_run("Tên nhân viên: ").bold = False
            run = p1.add_run(thong_tin.TenNV)
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) 

            # Thông tin Khách hàng
            p2 = document.add_paragraph()
            p2.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POS, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p2.add_run("Mã khách hàng: ").bold = False
            run = p2.add_run(thong_tin.MaKH)
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) 
            p2.add_run('\t')
            p2.add_run("Tên khách hàng: ").bold = False
            run = p2.add_run(thong_tin.TenKH)
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) 

            # Thông tin Ngày bán / Trạng thái
            p3 = document.add_paragraph()
            p3.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POS, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p3.add_run("Ngày bán: ").bold = False
            run = p3.add_run(ngay_ban_date.strftime("%d/%m/%Y"))
            run.bold = True
            run.font.color.rgb = RGBColor(0x43, 0xA0, 0x47) 
            p3.add_run('\t')
            p3.add_run("Trạng thái: ").bold = False
            run = p3.add_run(thong_tin.TrangThai)
            run.bold = True
            run.font.color.rgb = RGBColor(0x43, 0xA0, 0x47) 
                    
            document.add_paragraph() # Dòng trống

            # --- Bảng chi tiết hàng hóa ---
            table = document.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            headers = ["Mã CTHD", "Mã Tivi", "Tên Tivi", "Số Lượng", "Đơn Giá", "Thành Tiền"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ghi dữ liệu chi tiết
            for item in chi_tiet_hang:
                row_cells = table.add_row().cells
                row_cells[0].text = item.MaCTHD
                row_cells[1].text = item.MaTivi
                row_cells[2].text = item.TenTivi
                row_cells[3].text = str(item.SoLuong)
                row_cells[4].text = f"{float(item.DonGia):,.0f} đ"
                row_cells[5].text = f"{float(item.ThanhTien):,.0f} đ"
                
                # Căn giữa Số Lượng, Đơn Giá, Thành Tiền
                for i in [3, 4, 5]:
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # --- Tổng tiền (Căn phải) ---
            document.add_paragraph() 
            tong_tien_para = document.add_paragraph()
            tong_tien_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
            
            tong_tien_para.add_run("Tổng tiền: ").bold = True
            run = tong_tien_para.add_run(tong_tien_str)
            run.bold = True
            run.font.size = Pt(12) 
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) 

            # --- KHU VỰC KÝ TÊN ---
            document.add_paragraph() 
            document.add_paragraph()
            
            # Tạo bảng 3 cột cho khu vực chữ ký
            signature_table = document.add_table(rows=2, cols=3)
            signature_table.autofit = False
            signature_table.style = None 

            # Hàng 1: Chức danh
            chuc_danh = ["Người lập hóa đơn", "Kế toán bán hàng", "Khách hàng"]
            for i, cd in enumerate(chuc_danh):
                cell = signature_table.rows[0].cells[i]
                cell.text = cd
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].bold = True
            
            # Hàng 2: (Ghi rõ họ tên)
            for i in range(3):
                cell = signature_table.rows[1].cells[i]
                # Thêm tên nhân viên lập phiếu vào ô đầu tiên
                ten_nguoi_lap = thong_tin.TenNV if i == 0 else " " 
                ten_khach_hang = thong_tin.TenKH if i == 2 else " "
                
                if i == 0:
                     cell.text = f"(Ký, ghi rõ họ tên) \n{ten_nguoi_lap}"
                elif i == 2:
                     cell.text = f"(Ký, ghi rõ họ tên) \n{ten_khach_hang}"
                else:
                     cell.text = "(Ký, ghi rõ họ tên)"
                     
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Lưu file
            default_filename = f"HoaDonBanHang_{ma_hd}.docx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                initialfile=default_filename,
                filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
            )
            if file_path:
                document.save(file_path)
                messagebox.showinfo("Thành công", f"Đã in hóa đơn {ma_hd} ra file Word thành công tại:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Lỗi File Word", "Không thể tạo hoặc lưu file Word:\n" + str(e))

    def XoaHoaDonVinhVien(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 hóa đơn để xóa.")
            return

        ma_hd = self.trHienThi.item(selected[0], "values")[0]
        trang_thai = self.trHienThi.item(selected[0], "values")[5]
        
        # LẤY NGÀY BÁN CỦA HÓA ĐƠN TỪ CSDL
        try:
            cursor_temp = self.conn.cursor()
            cursor_temp.execute("SELECT NgayBan FROM HOADONBAN WHERE MaHD = ?", (ma_hd,))
            ngay_ban_db = cursor_temp.fetchone()[0]
            cursor_temp.close()
            
            # Tính toán xem hóa đơn đã tồn tại trên 5 năm chưa
            ngay_gioi_han = datetime.now().date().replace(year=datetime.now().year - 5)
            # Kiểm tra: NgayBan có nhỏ hơn Ngày giới hạn 5 năm KHÔNG?
            da_hon_5_nam = ngay_ban_db < ngay_gioi_han 
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể kiểm tra ngày bán của hóa đơn:\n{str(e)}")
            return
            
        is_allowed_to_delete = False
        warning_message = ""
        
        if trang_thai == "Đã hủy":
            is_allowed_to_delete = True
            warning_message = f"Bạn có chắc chắn muốn xóa vĩnh viễn hóa đơn ĐÃ HỦY {ma_hd} không?"
            
        elif trang_thai == "Đã thanh toán" and da_hon_5_nam:
            is_allowed_to_delete = True
            warning_message = f"Hóa đơn ĐÃ THANH TOÁN {ma_hd} này đã được lập hơn 5 năm ({ngay_ban_db.strftime('%d/%m/%Y')}). Bạn có chắc chắn muốn XÓA VĨNH VIỄN không?"
        
        if not is_allowed_to_delete:
            messagebox.showwarning("Cảnh báo", "Chỉ có thể xóa vĩnh viễn hóa đơn ĐÃ HỦY, hoặc hóa đơn ĐÃ THANH TOÁN có thời gian lưu trữ trên 5 năm.")
            return

        traloi = messagebox.askyesno("Xác nhận XÓA VĨNH VIỄN", warning_message)
        
        if traloi:
            try:
                cursor = self.conn.cursor()
                
                # Xóa các bản ghi liên quan (Bảo Hành)
                cursor.execute("SELECT MaCTHD FROM ChiTietHoaDon WHERE MaHD = ?", (ma_hd,))
                cthd_list = [row[0] for row in cursor.fetchall()]
                if cthd_list:
                    placeholders = ', '.join(['?'] * len(cthd_list))
                    cursor.execute(f"DELETE FROM BaoHanh WHERE MaCTHD IN ({placeholders})", cthd_list)

                # Xóa Chi Tiết Hóa Đơn
                cursor.execute("DELETE FROM ChiTietHoaDon WHERE MaHD = ?", (ma_hd,))
                
                # Xóa Hóa Đơn Bán
                cursor.execute("DELETE FROM HoaDonBan WHERE MaHD = ?", (ma_hd,))
                
                self.conn.commit()
                cursor.close()
                messagebox.showinfo("Thành công", f"Hóa đơn {ma_hd} đã được xóa vĩnh viễn khỏi hệ thống!")
                
                self.load_hoa_don()
                
            except Exception as e:
                self.conn.rollback()
                messagebox.showerror("Lỗi", "Đã xảy ra lỗi khi xóa hóa đơn:\n" + str(e))