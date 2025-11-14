import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
import pyodbc
from docx import Document 
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

# === TAB PHIẾU NHẬP HÀNG ===
class tabPhieuNhapHang(tk.Frame):
    def __init__(self, parent, conn):
        super().__init__(parent, bg="white")

        # === CHUỖI KẾT NỐI ===
        self.conn = conn

        # === KHUNG TÌM KIẾM ===
        frame_search = tk.Frame(self, bg="#E3F2FD", padx=10, pady=10)
        frame_search.pack(fill="x", padx=20, pady=5)

        tk.Label(frame_search, text="🔍 Tìm kiếm:", font=("Segoe UI", 10), bg="#E3F2FD").pack(side="left", padx=5)
        self.txt_timkiem = tk.Entry(frame_search, font=("Segoe UI", 10), width=20)
        self.txt_timkiem.pack(side="left", padx=5)

        self.search_option = tk.StringVar(value="mapn")
        tk.Radiobutton(frame_search, text="Theo mã phiếu nhập", variable=self.search_option, value="mapn", bg="#E3F2FD", font=("Segoe UI", 10)).pack(side="left", padx=10)
        tk.Radiobutton(frame_search, text="Theo mã nhà cung cấp", variable=self.search_option, value="mancc", bg="#E3F2FD", font=("Segoe UI", 10)).pack(side="left" , padx=10)
        tk.Radiobutton(frame_search, text="Theo trạng thái phiếu", variable=self.search_option, value="trangthai", bg="#E3F2FD", font=("Segoe UI", 10)).pack(side="left", padx=10)
        tk.Button(frame_search, text="Tìm", font=("Segoe UI", 10, "bold"), command=self.TimKiem, bg="#1565C0", fg="white", bd=0, padx=10, pady=5).pack(side="left", padx=10)
        tk.Button(frame_search, text="Hủy", font=("Segoe UI", 10, "bold"), command=self.HuyTimKiem, bg="#E53935", fg="white", bd=0, padx=10, pady=5).pack(side="left", padx=10)

        # === BẢNG DANH SÁCH PHIẾU NHẬP HÀNG ===
        frame_table = tk.Frame(self, bg="white")
        frame_table.pack(fill="both", expand=True, padx=20, pady=10)

        columns = ("MaPhieuNhap", "NgayNhap", "MaNV", "MaNCC", "TongTien", "TrangThai")
        
        # --- Tạo Scrollbar ---
        scroll_y = ttk.Scrollbar(frame_table, orient="vertical")
        scroll_x = ttk.Scrollbar(frame_table, orient="horizontal")

        self.trHienThi = ttk.Treeview( frame_table, show="headings",  columns=columns, height=12, yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)

        # --- Gắn Scrollbar ---
        scroll_y.config(command=self.trHienThi.yview)
        scroll_x.config(command=self.trHienThi.xview)

        # --- Bố trí Scrollbar ---
        scroll_y.pack(side="right", fill="y")
        scroll_x.pack(side="bottom", fill="x")
        self.trHienThi.pack(fill="both", expand=True)

        self.trHienThi.heading("MaPhieuNhap", text="Mã phiếu nhập")
        self.trHienThi.heading("NgayNhap", text="Ngày nhập")
        self.trHienThi.heading("MaNV", text="Mã nhân viên")
        self.trHienThi.heading("MaNCC", text="Mã nhà cung cấp")
        self.trHienThi.heading("TongTien", text="Tổng tiền")
        self.trHienThi.heading("TrangThai", text="Trạng thái")

        self.trHienThi.column("MaPhieuNhap", width=100,)
        self.trHienThi.column("NgayNhap", width=100, anchor="center")
        self.trHienThi.column("MaNV", width=100,)
        self.trHienThi.column("MaNCC", width=100,)
        self.trHienThi.column("TongTien", width=100, anchor="center")
        self.trHienThi.column("TrangThai", width=100)

        # Thêm style cho Treeview
        style = ttk.Style()
        style.configure("Treeview.Heading", font=("Segoe UI", 11, "bold"))
        style.configure("Treeview", font=("Segoe UI", 10), rowheight=28)

        # ==== Nút thao tác ====
        frame_btn = tk.Frame(self, bg="white")
        frame_btn.pack(pady=10)

        tk.Button(frame_btn, text="🧐 Xem chi tiết", bg="#EC9428", fg="white", font=("Segoe UI", 10, "bold"), command=self.XemChiTiet, padx=15, pady=5, bd=0).pack(side="left", padx=5)
        tk.Button(frame_btn, text="✅ Duyệt phiếu nhập hàng", bg="#43A047", fg="white",  font=("Segoe UI", 10, "bold"), command=self.DuyetPhieuNhapHang, padx=15, pady=5, bd=0).pack(side="left", padx=5)
        tk.Button(frame_btn, text="🗑 Hủy phiếu nhập hàng", bg="#E53935", fg="white",  font=("Segoe UI", 10, "bold"), command=self.HuyPhieuNhapHang, padx=15, pady=5, bd=0).pack(side="left", padx=5)
        tk.Button(frame_btn, text="🔄 Làm mới", bg="#1E88E5", fg="white", font=("Segoe UI", 10, "bold"), command=self.load_phieu_nhap, padx=15, pady=5, bd=0).pack(side="left", padx=5)
        tk.Button(frame_btn, text="🗑️ Xóa", bg="#B71C1C", fg="white", font=("Segoe UI", 10, "bold"), command=self.XoaPhieuNhapVinhVien, padx=15, pady=5, bd=0).pack(side="left", padx=5)
        tk.Button(frame_btn, text="🖨️  In", bg="#E51E9C", fg="white", font=("Segoe UI", 10, "bold"), command=self.InPhieuNhapHang, padx=15, pady=5, bd=0).pack(side="left", padx=5)

        # === TẢI DỮ LIỆU pHIẾU NHẬP ===
        self.load_phieu_nhap()

    def load_phieu_nhap(self):  
        try:
            self.trHienThi.delete(*self.trHienThi.get_children())
            cursor = self.conn.cursor()
            cursor.execute("SELECT MaPhieuNhap, NgayNhap, MaNV, MaNCC, TongTien, TrangThai FROM PHIEUNHAPHANG")

            for row in cursor.fetchall():
                ngay_nhap = datetime.strptime(str(row.NgayNhap).split(" ")[0], "%Y-%m-%d")

                formatted_row = (
                    row.MaPhieuNhap, ngay_nhap.strftime("%d/%m/%y"), row.MaNV, row.MaNCC,
                    f"{float(row.TongTien):,.0f}" if row.TongTien else "0", row.TrangThai)
                
                self.trHienThi.insert("", tk.END, values=formatted_row)
                    
        except Exception as e:
            messagebox.showerror("Lỗi", "Không thể tải dữ liệu hóa đơn:\n + " + str(e))

    # === Hàm xem chi tiết phiếu nhập ===
    def XemChiTiet(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 phiếu nhập để xem chi tiết.")
            return

        ma_phieu = self.trHienThi.item(selected[0], "values")[0]

        try:
            chitietphieunhap= tk.Toplevel(self)
            chitietphieunhap.title("Chi tiết phiếu nhập: " + ma_phieu)
            chitietphieunhap.geometry("700x800")
            chitietphieunhap.resizable(False, False)
            chitietphieunhap.configure(bg="white")

            tk.Label(chitietphieunhap, text="Chi tiết phiếu nhập hàng: " + ma_phieu, font=("Segoe UI", 12, "bold"), bg="white", fg="#0D47A1").pack(pady=10)
            
            cursor = self.conn.cursor()
            cursor.execute("""SELECT nv.MaNV, nv.TenNV, ncc.MaNCC, ncc.TenNCC, pnh.NgayNhap, pnh.TongTien, pnh.TrangThai
                              FROM PHIEUNHAPHANG pnh JOIN NHANVIEN nv ON pnh.MaNV = nv.MaNV
                              JOIN NHACUNGCAP ncc ON pnh.MaNCC = ncc.MaNCC
                              WHERE pnh.MaPhieuNhap = ?""", (ma_phieu,))
            thong_tin = cursor.fetchone()

            Frame_thongtin = tk.Frame(chitietphieunhap, bg="white")
            Frame_thongtin.pack(pady=5)
            tk.Label(Frame_thongtin, text="Mã nhân viên:", font=("Segoe UI", 10), bg="white").grid(row=0, column=0, padx=10, sticky="w")
            tk.Label(Frame_thongtin, text=thong_tin.MaNV, font=("Segoe UI", 10, "bold"), bg="white", fg="#1565C0").grid(row=0, column=1, padx=10, sticky="w")

            tk.Label(Frame_thongtin, text="Tên nhân viên:", font=("Segoe UI", 10), bg="white").grid(row=0, column=2, padx=10, sticky="w")
            tk.Label(Frame_thongtin, text=thong_tin.TenNV, font=("Segoe UI", 10, "bold"), bg="white", fg="#1565C0").grid(row=0, column=3, padx=10, sticky="w")

            tk.Label(Frame_thongtin, text="Mã nhà cung cấp:", font=("Segoe UI", 10), bg="white").grid(row=1, column=0, padx=10, sticky="w")
            tk.Label(Frame_thongtin, text=thong_tin.MaNCC, font=("Segoe UI", 10, "bold"), bg="white", fg="#1565C0").grid(row=1, column=1, padx=10, sticky="w")

            tk.Label(Frame_thongtin, text="Tên nhà cung cấp:", font=("Segoe UI", 10), bg="white").grid(row=1, column=2, padx=10, sticky="w")
            tk.Label(Frame_thongtin, text=thong_tin.TenNCC, font=("Segoe UI", 10, "bold"), bg="white", fg="#1565C0").grid(row=1, column=3, padx=10, sticky="w")

            tk.Label(Frame_thongtin, text="Ngày nhập:", font=("Segoe UI", 10), bg="white").grid(row=2, column=0, padx=10, sticky="w")
            
            ngay_nhap = datetime.strptime(str(thong_tin.NgayNhap).split(" ")[0], "%Y-%m-%d")
            tk.Label(Frame_thongtin, text=ngay_nhap.strftime("%d/%m/%Y"), font=("Segoe UI", 10, "bold"), bg="white", fg="#43A047").grid(row=2, column=1, padx=10, sticky="w")

            tk.Label(Frame_thongtin, text="Trạng thái:", font=("Segoe UI", 10), bg="white").grid(row=2, column=2, padx=10, sticky="w")
            tk.Label(Frame_thongtin, text=thong_tin.TrangThai, font=("Segoe UI", 10, "bold"), bg="white", fg="#43A047").grid(row=2, column=3, padx=10, sticky="w")
            
            columns = ("MaTivi", "TenTivi", "SoLuong", "GiaNhap", "ThanhTien")
            tree = ttk.Treeview(chitietphieunhap, columns=columns, show="headings", height=10)
            tree.pack(fill="both", expand=True, padx=15, pady=10)

            tree.heading("MaTivi", text="Mã Tivi")
            tree.heading("TenTivi", text="Tên Tivi")
            tree.heading("SoLuong", text="Số Lượng")
            tree.heading("GiaNhap", text="Giá Nhập")
            tree.heading("ThanhTien", text="Thành Tiền")

            tree.column("MaTivi", width=100)
            tree.column("TenTivi", width=200)
            tree.column("SoLuong", width=100, anchor="center")
            tree.column("GiaNhap", width=100, anchor="center")
            tree.column("ThanhTien", width=100, anchor="center")

            tk.Label(chitietphieunhap, text="Tổng tiền:", font=("Segoe UI", 10, "bold"), bg="white").pack(side="left", padx=20)
            tk.Label(chitietphieunhap, text=f"{float(thong_tin.TongTien):,.0f} đ", font=("Segoe UI", 10, "bold"), bg="white", fg="red").pack(side="right", padx=20)
            
            cursor.execute("""
                SELECT ctpn.MaTivi, tv.TenTivi, ctpn.SoLuong, ctpn.GiaNhap, ctpn.ThanhTien
                FROM CHITIETPHIEUNHAP ctpn
                JOIN TIVI tv ON ctpn.MaTivi = tv.MaTivi
                WHERE ctpn.MaPhieuNhap = ?
            """, (ma_phieu,))
            rows = cursor.fetchall()

            if not rows:
                messagebox.showinfo("Thông báo", "Không có chi tiết cho phiếu " + ma_phieu)
                chitietphieunhap.destroy()
                return

            for r in rows:
                tree.insert("", tk.END, values=(r.MaTivi, r.TenTivi,  r.SoLuong, f"{float(r.GiaNhap):,.0f} đ", f"{float(r.ThanhTien):,.0f} đ"))

            cursor.close()
        except Exception as e:
            messagebox.showerror("Lỗi", "Không thể xem chi tiết phiếu nhập:\n" + str(e))

    def DuyetPhieuNhapHang(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 phiếu nhập để duyệt.")
            return
        ma_phieu = self.trHienThi.item(selected[0], "values")[0]
        trang_thai = self.trHienThi.item(selected[0], "values")[5]

        if trang_thai != "Đợi duyệt":
            messagebox.showwarning("Thông báo", "Chỉ có thể duyệt phiếu nhập hàng chưa được duyệt.")
            return
        
        traloi = messagebox.askyesno("Xác nhận", "Bạn có chắc chắn muốn duyệt phiếu nhập hàng " + ma_phieu + " không?")
        if traloi:
            try:
                cursor = self.conn.cursor()
                # Cập nhật duyệt hàng
                cursor.execute("""
                    UPDATE PHIEUNHAPHANG
                    SET TrangThai = N'Đã duyệt'
                    WHERE MaPhieuNhap = ?
                """, (ma_phieu,))

                # Lấy chi tiết phiếu nhập hàng để cập nhật tồn kho
                cursor.execute("""
                    SELECT MaTivi, SoLuong
                    FROM CHITIETPHIEUNHAP
                    WHERE MaPhieuNhap = ?
                """, (ma_phieu,))

                chitiet = cursor.fetchall()

                for item in chitiet:
                    ma_tivi = item.MaTivi
                    so_luong = item.SoLuong

                    # Cập nhật số lượng tồn kho của tivi
                    cursor.execute("""
                        UPDATE TIVI
                        SET SoLuongTon = SoLuongTon + ?
                        WHERE MaTivi = ? """, (so_luong, ma_tivi))
                
                self.conn.commit()
                cursor.close()
                messagebox.showinfo("Thành công", "Phiếu nhập hàng đã được duyệt thành công!")

                # Làm mới lại danh sách phiếu nhập hàng
                self.load_phieu_nhap()

            except Exception as e:
                messagebox.showerror("Lỗi", "Đã xảy ra lỗi khi duyệt phiếu nhập hàng:\n" + str(e))

    def HuyPhieuNhapHang(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 phiếu nhập để hủy")
            return
        
        ma_phieu = self.trHienThi.item(selected[0], "values")[0]
        trang_thai = self.trHienThi.item(selected[0], "values")[5]

        if trang_thai != "Đợi duyệt":
            messagebox.showwarning("Thông báo", "Chỉ có thể hủy phiếu nhập hàng chưa được duyệt.")
            return
        
        traloi = messagebox.askyesno("Xác nhận", "Bạn có chắc chắn muốn hủy phiếu nhập hàng " + ma_phieu + " không?")
        if traloi:
            try:
                cursor = self.conn.cursor()
                cursor.execute("""
                    UPDATE PHIEUNHAPHANG
                    SET TrangThai = N'Đã hủy'
                    WHERE MaPhieuNhap = ?""", (ma_phieu,))
                
                self.conn.commit()
                cursor.close()
                messagebox.showinfo("Thành công", "Phiếu nhập hàng đã được hủy thành công!")
                self.load_phieu_nhap()

            except Exception as e:
                messagebox.showerror("Lỗi", "Đã xảy ra lỗi khi hủy phiếu nhập hàng:\n" + str(e))

    def TimKiem(self):
        if self.txt_timkiem.get().strip() == "":
            messagebox.showwarning("Thông báo", "Vui lòng nhập từ khóa tìm kiếm.")
            return
        
        timkkiem = self.search_option.get().strip()
        self.trHienThi.delete(*self.trHienThi.get_children())

        try:
            cursor = self.conn.cursor()

            if timkkiem == "mapn":
                keyword = self.txt_timkiem.get()
                cursor.execute("""
                    SELECT MaPhieuNhap, NgayNhap, MaNV, MaNCC, TongTien, TrangThai
                    FROM PHIEUNHAPHANG
                    WHERE MaPhieuNhap = ?""", (keyword,))
            
            elif timkkiem == "mancc":
                keyword = self.txt_timkiem.get()
                cursor.execute("""
                    SELECT MaPhieuNhap, NgayNhap, MaNV, MaNCC, TongTien, TrangThai
                    FROM PHIEUNHAPHANG
                    WHERE MaNCC = ?""", (keyword,))
            
            elif timkkiem == "trangthai":
                keyword = self.txt_timkiem.get()
                cursor.execute("""
                    SELECT MaPhieuNhap, NgayNhap, MaNV, MaNCC, TongTien, TrangThai
                    FROM PHIEUNHAPHANG
                    WHERE TrangThai = ?""", (keyword,))
                
            rows = cursor.fetchall()
            for row in rows:
                ngay_nhap = datetime.strptime(str(row.NgayNhap).split(" ")[0], "%Y-%m-%d")

                formatted_row = (
                    row.MaPhieuNhap, ngay_nhap.strftime("%d/%m/%y"), row.MaNV, row.MaNCC,
                    f"{float(row.TongTien):,.0f}" if row.TongTien else "0", row.TrangThai
                )
                self.trHienThi.insert("", tk.END, values=formatted_row)
            cursor.close()

        except Exception as e:
            messagebox.showerror("Lỗi", "Không thể tìm kiếm phiếu nhập hàng:\n" + str(e))

    def HuyTimKiem(self):
        # Xóa txt_timkiem nếu có
        self.load_phieu_nhap()
        self.txt_timkiem.delete(0, tk.END)

    def XoaPhieuNhapVinhVien(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 phiếu nhập để xóa.")
            return

        ma_phieu = self.trHienThi.item(selected[0], "values")[0]
        trang_thai = self.trHienThi.item(selected[0], "values")[5]
        
        # LẤY NGÀY NHẬP CỦA PHIẾU ĐƯỢC CHỌN TỪ CSDL
        try:
            cursor_temp = self.conn.cursor()
            cursor_temp.execute("SELECT NgayNhap FROM PHIEUNHAPHANG WHERE MaPhieuNhap = ?", (ma_phieu,))
            ngay_nhap_db = cursor_temp.fetchone()[0]
            cursor_temp.close()
            
            # Tính toán xem phiếu đã tồn tại trên 5 năm chưa
            ngay_gioi_han = datetime.now().date().replace(year=datetime.now().year - 5)
            da_hon_5_nam = ngay_nhap_db < ngay_gioi_han 
            
        except Exception as e:
            messagebox.showerror("Lỗi", "Không thể kiểm tra ngày nhập của phiếu: " + str(e))
            return
        
        is_allowed_to_delete = False
        warning_message = ""
        
        if trang_thai == "Đã hủy":
            is_allowed_to_delete = True
            warning_message = f"Bạn có chắc chắn muốn xóa vĩnh viễn phiếu nhập ĐÃ HỦY {ma_phieu} không?"
            
        elif trang_thai == "Đã duyệt" and da_hon_5_nam:
            is_allowed_to_delete = True
            warning_message = f"Phiếu nhập ĐÃ DUYỆT {ma_phieu} này đã được lưu trữ hơn 5 năm ({ngay_nhap_db.strftime('%d/%m/%Y')}). Bạn có chắc chắn muốn XÓA VĨNH VIỄN không?"
        
        if not is_allowed_to_delete:
            messagebox.showwarning("Cảnh báo", "Chỉ có thể xóa vĩnh viễn phiếu nhập ĐÃ HỦY, hoặc phiếu ĐÃ DUYỆT có thời gian lưu trữ trên 5 năm.")
            return

        traloi = messagebox.askyesno("Xác nhận XÓA VĨNH VIỄN", warning_message)
        
        if traloi:
            try:
                # Bắt đầu xóa CSDL
                cursor = self.conn.cursor()
                
                # Xóa Chi Tiết Phiếu Nhập trước
                cursor.execute("DELETE FROM ChiTietPhieuNhap WHERE MaPhieuNhap = ?", (ma_phieu,))
                
                # Xóa Phiếu Nhập Hàng
                cursor.execute("DELETE FROM PhieuNhapHang WHERE MaPhieuNhap = ?", (ma_phieu,))
                
                self.conn.commit()
                cursor.close()
                messagebox.showinfo("Thành công", f"Phiếu nhập hàng {ma_phieu} đã được xóa vĩnh viễn khỏi hệ thống!")
                
                self.load_phieu_nhap()         
            except Exception as e:
                messagebox.showerror("Lỗi CSDL", f"Đã xảy ra lỗi khi xóa phiếu nhập: "  + str(e))

    # === HÀM IN PHIẾU NHẬP HÀNG RA WORD ===
    def InPhieuNhapHang(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 phiếu nhập để in.")
            return

        ma_phieu = self.trHienThi.item(selected[0], "values")[0]
        
        # --- Lấy dữ liệu chi tiết ---
        try:
            cursor = self.conn.cursor()
            
            # 1. Lấy thông tin chung của phiếu nhập (Giữ nguyên)
            cursor.execute("""SELECT pnh.MaPhieuNhap, nv.MaNV, nv.TenNV, ncc.MaNCC, ncc.TenNCC, pnh.NgayNhap, pnh.TongTien, pnh.TrangThai
                              FROM PHIEUNHAPHANG pnh 
                              JOIN NHANVIEN nv ON pnh.MaNV = nv.MaNV
                              JOIN NHACUNGCAP ncc ON pnh.MaNCC = ncc.MaNCC
                              WHERE pnh.MaPhieuNhap = ?""", (ma_phieu,))
            thong_tin = cursor.fetchone()

            if not thong_tin:
                messagebox.showerror("Lỗi", "Không tìm thấy thông tin phiếu nhập " + ma_phieu)
                return

            # 2. Lấy chi tiết các mặt hàng (Giữ nguyên)
            cursor.execute("""
                SELECT ctpn.MaTivi, tv.TenTivi, ctpn.SoLuong, ctpn.GiaNhap, ctpn.ThanhTien
                FROM CHITIETPHIEUNHAP ctpn
                JOIN TIVI tv ON ctpn.MaTivi = tv.MaTivi
                WHERE ctpn.MaPhieuNhap = ?
            """, (ma_phieu,))
            chi_tiet_hang = cursor.fetchall()
            cursor.close()

            if not chi_tiet_hang:
                messagebox.showinfo("Thông báo", "Phiếu nhập " + ma_phieu + " không có chi tiết hàng hóa.")
                return

        except Exception as e:
            messagebox.showerror("Lỗi CSDL", "Không thể lấy dữ liệu để in:\n" + str(e))
            return
            
        # --- Tạo File Word ---
        try:
            document = Document()
            
            # Đặt font và size mặc định cho document
            style = document.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(11)

            # Dữ liệu định dạng
            ngay_nhap_date = datetime.strptime(str(thong_tin.NgayNhap).split(" ")[0], "%Y-%m-%d")
            ngay_nhap_str = ngay_nhap_date.strftime("%d/%m/%Y")
            tong_tien_str = f"{float(thong_tin.TongTien):,.0f} đ"
            
            # Tiêu đề chính (Căn giữa)
            heading = document.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.add_run(f"PHIẾU NHẬP HÀNG SỐ: {ma_phieu}")
            run.bold = True
            run.font.size = Pt(16)
            document.add_paragraph() # Dòng trống
            
            # --- Bố trí Thông tin chung  ---
            
            # Thiết lập Tab Stop: Căn lề trái tại vị trí 8.0 cm 
            TAB_STOP_POS = Cm(5.0) 
            
            # Thông tin Nhân viên
            p1 = document.add_paragraph()
            p1.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POS, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p1.add_run("Mã nhân viên: ").bold = False
            run = p1.add_run(thong_tin.MaNV)
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) 
            p1.add_run('\t') 
            p1.add_run("Tên nhân viên: ").bold = False
            run = p1.add_run(thong_tin.TenNV)
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) 

            # Thông tin Nhà cung cấp
            p2 = document.add_paragraph()
            p2.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POS, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p2.add_run("Mã nhà cung cấp: ").bold = False
            run = p2.add_run(thong_tin.MaNCC)
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) 
            p2.add_run('\t')
            p2.add_run("Tên nhà cung cấp: ").bold = False
            run = p2.add_run(thong_tin.TenNCC)
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) 

            # Thông tin Ngày nhập / Trạng thái
            p3 = document.add_paragraph()
            p3.paragraph_format.tab_stops.add_tab_stop(TAB_STOP_POS, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p3.add_run("Ngày nhập: ").bold = False
            run = p3.add_run(ngay_nhap_str)
            run.bold = True
            run.font.color.rgb = RGBColor(0x43, 0xA0, 0x47) 
            p3.add_run('\t')
            p3.add_run("Trạng thái: ").bold = False
            run = p3.add_run(thong_tin.TrangThai)
            run.bold = True
            run.font.color.rgb = RGBColor(0x43, 0xA0, 0x47) 
                    
            document.add_paragraph() # Dòng trống

            # --- Bảng chi tiết hàng hóa (Giữ nguyên) ---
            table = document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            # ... (Phần headers và ghi dữ liệu chi tiết giữ nguyên) ...
            
            hdr_cells = table.rows[0].cells
            headers = ["Mã Tivi", "Tên Tivi", "Số Lượng", "Giá Nhập", "Thành Tiền"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for item in chi_tiet_hang:
                row_cells = table.add_row().cells
                row_cells[0].text = item.MaTivi
                row_cells[1].text = item.TenTivi
                row_cells[2].text = str(item.SoLuong)
                row_cells[3].text = f"{float(item.GiaNhap):,.0f} đ"
                row_cells[4].text = f"{float(item.ThanhTien):,.0f} đ"
                for i in [2, 3, 4]:
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # --- Tổng tiền (Căn phải) ---
            document.add_paragraph() 
            tong_tien_para = document.add_paragraph()
            tong_tien_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
            
            tong_tien_para.add_run("Tổng tiền: ").bold = True
            run = tong_tien_para.add_run(tong_tien_str)
            run.bold = True
            run.font.size = Pt(12) 
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) 

            # --- KHU VỰC KÝ TÊN ---
            document.add_paragraph() 
            document.add_paragraph()
            
            # Tạo bảng 3 cột cho khu vực chữ ký
            signature_table = document.add_table(rows=2, cols=3)
            signature_table.autofit = False
            signature_table.style = None # Bỏ style lưới để không có viền

            # Hàng 1: Chức danh
            chuc_danh = ["Người lập phiếu", "Thủ kho", "Kế toán trưởng"]
            for i, cd in enumerate(chuc_danh):
                cell = signature_table.rows[0].cells[i]
                cell.text = cd
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].bold = True
            
            # Hàng 2: (Ghi rõ họ tên)
            for i in range(3):
                cell = signature_table.rows[1].cells[i]
                # Thêm tên nhân viên lập phiếu vào ô đầu tiên
                ten_nguoi_lap = thong_tin.TenNV if i == 0 else " " # Thêm tên người lập vào ô đầu tiên
                cell.text = f"(Ký, ghi rõ họ tên) \n{ten_nguoi_lap}"
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Lưu file
            default_filename = f"PhieuNhapHang_{ma_phieu}.docx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                initialfile=default_filename,
                filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
            )

            if file_path:
                document.save(file_path)
                messagebox.showinfo("Thành công", f"Đã in phiếu nhập hàng {ma_phieu} ra file Word thành công tại:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Lỗi File Word", "Không thể tạo hoặc lưu file Word:\n" + str(e))
            