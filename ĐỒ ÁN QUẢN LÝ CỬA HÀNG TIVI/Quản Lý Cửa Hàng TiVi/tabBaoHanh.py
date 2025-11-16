import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkcalendar import DateEntry
from datetime import date, datetime, timedelta
import pyodbc
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER


class tabBaoHanh(tk.Frame):
    def __init__(self, parent, conn, user):
        super().__init__(parent, bg="white")

        self.conn = conn
        self.cursor = conn.cursor()

        self.user = user

        self.ds_them = []
        self.ds_sua = []
        self.ds_xoa = []

        # === KHUNG TÌM KIẾM ===
        frame_search = tk.Frame(self, bg="#E3F2FD", padx=10, pady=10)
        frame_search.pack(fill="x", padx=20, pady=5)

        tk.Label(frame_search, text="Tìm kiếm:", font=("Segoe UI", 10), bg="#E3F2FD").pack(side="left", padx=5)
        self.txt_timkiem = tk.Entry(frame_search, font=("Segoe UI", 10), width=30, bg="white")
        self.txt_timkiem.pack(side="left", padx=5)
        self.txt_timkiem.bind("<Return>", lambda e: self.timkiem())

        self.search_option = tk.StringVar(master=self, value="mabh")
        tk.Radiobutton(frame_search, text="Mã bảo hành", variable=self.search_option, value="mabh", bg="#E3F2FD", font=("Segoe UI", 10)).pack(side="left", padx=5)
        tk.Radiobutton(frame_search, text="Mã chi tiết hóa đơn", variable=self.search_option, value="macthd", bg="#E3F2FD", font=("Segoe UI", 10)).pack(side="left", padx=5)
        tk.Radiobutton(frame_search, text="Mã hóa đơn", variable=self.search_option, value="mahd", bg="#E3F2FD", font=("Segoe UI", 10)).pack(side="left", padx=5)

        tk.Button(frame_search, text="Tìm", font=("Segoe UI", 10, "bold"), bg="#1565C0", fg="white", bd=0, padx=10, pady=5, command=self.timkiem).pack(side="left", padx=10)
        tk.Button(frame_search, text="Hủy", font=("Segoe UI", 10, "bold"), bg="#E53935", fg="white", bd=0, padx=10, pady=5, command=self.huy).pack(side="left", padx=10)

        # === KHUNG THÔNG TIN ===
        frame_form = tk.LabelFrame(self, text="Thông tin Bảo hành", bg="white", font=("Segoe UI", 12, "bold"), fg="#0D47A1", padx=10, pady=10)
        frame_form.pack(fill="x", padx=20, pady=10)

        tk.Label(frame_form, text="Mã BH:", bg="white", font=("Segoe UI", 10)).grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.txt_mabh = ttk.Entry(frame_form, width=34)
        self.txt_mabh.grid(row=0, column=1, padx=5, pady=5)

        tk.Label(frame_form, text="Mã CTHD:", bg="white", font=("Segoe UI", 10)).grid(row=0, column=2, sticky="w", padx=5, pady=5)
        self.cb_macthd = ttk.Combobox(frame_form, width=22, font=("Segoe UI", 10), state="readonly")
        self.cb_macthd.grid(row=0, column=3, padx=5, pady=5)
        self.cb_macthd.bind("<<ComboboxSelected>>", lambda e: self.capnhat_mahd_theo_cthd())

        tk.Label(frame_form, text="Mã HD:", bg="white", font=("Segoe UI", 10)).grid(row=0, column=4, sticky="w", padx=5, pady=5)
        self.cb_mahd = ttk.Combobox(frame_form, width=22, font=("Segoe UI", 10), state="readonly")
        self.cb_mahd.grid(row=0, column=5, padx=5, pady=5)

        tk.Label(frame_form, text="Thời gian (tháng):", bg="white", font=("Segoe UI", 10)).grid(row=1, column=0, sticky="w", padx=5, pady=5)
        self.txt_thoigian = ttk.Entry(frame_form, width=34)
        self.txt_thoigian.grid(row=1, column=1, padx=5, pady=5)

        tk.Label(frame_form, text="Điều kiện:", bg="white", font=("Segoe UI", 10)).grid(row=1, column=2, sticky="w", padx=5, pady=5)
        self.txt_dieukien = ttk.Entry(frame_form, width=45)
        self.txt_dieukien.grid(row=1, column=3, columnspan=3, padx=5, pady=5, sticky="we")

        tk.Label(frame_form, text="Ngày BH:", bg="white", font=("Segoe UI", 10)).grid(row=2, column=0, sticky="w", padx=5, pady=5)
        self.date_ngaybaohanh = DateEntry(frame_form, width=29, font=("Segoe UI", 10), date_pattern="dd/mm/yyyy")
        self.date_ngaybaohanh.grid(row=2, column=1, padx=5, pady=5)

        tk.Label(frame_form, text="Trạng thái:", bg="white", font=("Segoe UI", 10)).grid(row=2, column=2, sticky="w", padx=5, pady=5)
        self.lbl_trangthai = tk.Label(frame_form, text="", bg="white", font=("Segoe UI", 10, "bold"))
        self.lbl_trangthai.grid(row=2, column=3, sticky="w", padx=5, pady=5)

        # === NÚT CHỨC NĂNG ===
        frame_buttons = tk.Frame(self, bg="white")
        frame_buttons.pack(pady=10)

        if self.user == "admin":
            tk.Button(frame_buttons, text="➕ Thêm", bg="#EBDA42", fg="white", font=("Segoe UI", 11, "bold"), padx=20, pady=5, bd=0, command=self.them).grid(row=0, column=0, padx=8)
            tk.Button(frame_buttons, text="✏️ Sửa", bg="#FB8C00", fg="white", font=("Segoe UI", 11, "bold"), padx=20, pady=5, bd=0, command=self.sua).grid(row=0, column=1, padx=8)
            tk.Button(frame_buttons, text="🗑️ Xóa", bg="#E53935", fg="white", font=("Segoe UI", 11, "bold"), padx=20, pady=5, bd=0, command=self.xoa).grid(row=0, column=2, padx=8)
            tk.Button(frame_buttons, text="🔄 Làm mới", bg="#1E88E5", fg="white", font=("Segoe UI", 11, "bold"), padx=20, pady=5, bd=0, command=self.lammoi).grid(row=0, column=3, padx=8)
            tk.Button(frame_buttons, text="💾 Lưu", bg="#43A047", fg="white", font=("Segoe UI", 11, "bold"), padx=20, pady=5, bd=0, command=self.luu).grid(row=0, column=4, padx=8)
        tk.Button(frame_buttons, text="🖨️ In bảo hành", bg="#E51E9C", fg="white", font=("Segoe UI", 11, "bold"), padx=15, command=self.InBaoHanh, pady=5, bd=0).grid(row=0, column=5, padx=8)

        # === BẢNG HIỂN THỊ ===
        frame_table = tk.Frame(self, bg="white")
        frame_table.pack(fill="both", expand=True, padx=20, pady=10)

        columns = ("MaBH", "MaCTHD", "MaHD", "ThoiGianBaoHanh", "DieuKien", "NgayBaoHanh", "TrangThai")

        scroll_y = ttk.Scrollbar(frame_table, orient="vertical")
        scroll_x = ttk.Scrollbar(frame_table, orient="horizontal")

        self.trHienThi = ttk.Treeview(frame_table, columns=columns, show="headings", height=14, yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)

        scroll_y.pack(side="right", fill="y")
        scroll_x.pack(side="bottom", fill="x")
        self.trHienThi.pack(fill="both", expand=True)

        self.trHienThi.heading("MaBH", text="Mã BH")
        self.trHienThi.heading("MaCTHD", text="Mã CTHD")
        self.trHienThi.heading("MaHD", text="Mã HD")
        self.trHienThi.heading("ThoiGianBaoHanh", text="Thời gian (tháng)")
        self.trHienThi.heading("DieuKien", text="Điều kiện")
        self.trHienThi.heading("NgayBaoHanh", text="Ngày BH")
        self.trHienThi.heading("TrangThai", text="Trạng thái")

        self.trHienThi.column("MaBH", width=100, anchor="center")
        self.trHienThi.column("MaCTHD", width=100, anchor="center")
        self.trHienThi.column("MaHD", width=100, anchor="center")
        self.trHienThi.column("ThoiGianBaoHanh", width=130, anchor="center")
        self.trHienThi.column("DieuKien", width=250, anchor="w")
        self.trHienThi.column("NgayBaoHanh", width=110, anchor="center")
        self.trHienThi.column("TrangThai", width=120, anchor="center")

        style = ttk.Style()
        style.configure("Treeview.Heading", font=("Segoe UI", 11, "bold"))
        style.configure("Treeview", font=("Segoe UI", 10), rowheight=30)

        self.trHienThi.bind("<<TreeviewSelect>>", self.chon_dong)

        self.hien_thi_du_lieu_cthd()
        self.hienthi_dulieu()

    def hien_thi_du_lieu_cthd(self):
        try:
            sql_query = """
                SELECT DISTINCT cthd.MaCTHD
                FROM ChiTietHoaDon cthd
                JOIN HoaDonBan hdb ON hdb.MaHD = cthd.MaHD
                WHERE hdb.TrangThai = N'Đã thanh toán'
            """
            params = []
            
            if self.user.lower() != "admin":
                sql_query += " AND hdb.MaNV = ?"
                params.append(self.user)
            
            sql_query += " ORDER BY cthd.MaCTHD"
            
            self.cursor.execute(sql_query, params)
            rows = self.cursor.fetchall()
            self.cb_macthd["values"] = [row.MaCTHD for row in rows]

            if rows:
                self.cb_macthd.current(0)
                self.capnhat_mahd_theo_cthd()

        except Exception as e:
            messagebox.showerror("Lỗi", f"Không tải được danh sách CTHD: {str(e)}")

    def capnhat_mahd_theo_cthd(self):
        macthd = self.cb_macthd.get()
        if not macthd:
            self.cb_mahd["values"] = []
            return
        try:
            self.cursor.execute(" SELECT MaHD FROM ChiTietHoaDon WHERE MaCTHD = ? ", (macthd,))
            row = self.cursor.fetchone()
            if row:
                self.cb_mahd["values"] = [row.MaHD]
                self.cb_mahd.set(row.MaHD)
            else:
                self.cb_mahd["values"] = []
                self.cb_mahd.set("")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không tải được Mã HD: {str(e)}")

    def hienthi_dulieu(self):
        for item in self.trHienThi.get_children():
            self.trHienThi.delete(item)

        try:
            sql_query = """
                SELECT bh.MaBH, bh.MaCTHD, cthd.MaHD, hdb.MaNV, bh.ThoiGianBaoHanh, bh.DieuKien, bh.NgayBaoHanh
                FROM BaoHanh bh
                JOIN ChiTietHoaDon cthd ON bh.MaCTHD = cthd.MaCTHD
                JOIN HoaDonBan hdb ON cthd.MaHD = hdb.MaHD"""
            params = []

            if self.user != "admin":
                sql_query += " WHERE hdb.MaNV = ?"
                params.append(self.user)
            
            sql_query += " ORDER BY bh.MaBH"
            
            self.cursor.execute(sql_query, params)
            rows = self.cursor.fetchall()

            for row in rows:
                ngay_bh = row.NgayBaoHanh.date() if hasattr(row.NgayBaoHanh, 'date') else date.fromisoformat(str(row.NgayBaoHanh).split()[0])
                
                # Tính Hạn/Hết Hạn dựa trên NgayBaoHanh trong CSDL
                ngay_het = ngay_bh + timedelta(days=row.ThoiGianBaoHanh * 30)
                trangthai_hienthi = "CÒN HẠN" if ngay_het >= date.today() else "HẾT HẠN"
                    
                # Màu sắc hiển thị
                tag = 'con_han' if trangthai_hienthi == 'CÒN HẠN' else 'het_han'

                self.trHienThi.insert("", "end", 
                    values=(row.MaBH, row.MaCTHD, row.MaHD, row.ThoiGianBaoHanh, row.DieuKien or "", self.chuyen_yyyy_sang_dd(row.NgayBaoHanh), trangthai_hienthi),
                    tags=(tag,))
            
        except Exception as e:
            messagebox.showerror("Lỗi", "Không thể tải dữ liệu: " + str(e))

    def chuyen_yyyy_sang_dd(self, ngay_db):
        if ngay_db is None:
            return ""
        ngay_str = str(ngay_db).strip().split()[0]
        if "-" in ngay_str and len(ngay_str.split("-")) == 3:
            y, m, d = ngay_str.split("-")
            return f"{d.zfill(2)}/{m.zfill(2)}/{y}"
        return ngay_str

    def chon_dong(self, event):
        selected = self.trHienThi.selection()
        if not selected:
            return

        values = self.trHienThi.item(selected[0])["values"]
        self.xoa_form()

        self.txt_mabh.insert(0, values[0])
        self.cb_macthd.set(values[1])
        self.capnhat_mahd_theo_cthd()
        self.txt_thoigian.insert(0, values[3])
        self.txt_dieukien.insert(0, values[4])

        d, m, y = map(int, values[5].split("/"))
        self.date_ngaybaohanh.set_date(date(y, m, d))

        self.lbl_trangthai.config(text="CÒN HẠN BẢO HÀNH" if values[6] == "CÒN HẠN" else "ĐÃ HẾT HẠN BẢO HÀNH", fg="green" if values[6] == "CÒN HẠN" else "red")

    def them(self):
        mabh = self.txt_mabh.get().strip()
        macthd = self.cb_macthd.get()
        mahd = self.cb_mahd.get()
        thoigian = self.txt_thoigian.get().strip()
        dieukien = self.txt_dieukien.get().strip()
        ngaybh = self.date_ngaybaohanh.get_date()

        if not all([mabh, macthd, thoigian]):
            messagebox.showwarning("Cảnh báo", "Vui lòng nhập đầy đủ thông tin bắt buộc!")
            return

        try:
            thoigian_int = int(thoigian)
            if thoigian_int <= 0:
                raise ValueError
        except:
            messagebox.showwarning("Cảnh báo", "Thời gian phải là số nguyên dương!")
            return

        for item in self.trHienThi.get_children():
            if self.trHienThi.item(item)["values"][0] == mabh:
                messagebox.showwarning("Cảnh báo", f"Mã BH '{mabh}' đã tồn tại!")
                return

        ngay_het = ngaybh + timedelta(days=thoigian_int * 30)
        trangthai = "CÒN HẠN" if ngay_het >= date.today() else "HẾT HẠN"

        self.trHienThi.insert("", "end", values=(mabh, macthd, mahd, thoigian_int, dieukien, self.chuyen_yyyy_sang_dd(ngaybh), trangthai))

        self.ds_them.append((mabh, macthd, thoigian_int, dieukien, ngaybh))
        self.xoa_form()
        messagebox.showinfo("Thành công", "Đã thêm! Nhấn 'Lưu' để lưu vào CSDL.")

    def sua(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Cảnh báo", "Chọn dòng cần sửa!")
            return

        mabh = self.txt_mabh.get().strip()
        macthd = self.cb_macthd.get()
        thoigian = self.txt_thoigian.get().strip()
        dieukien = self.txt_dieukien.get().strip()
        ngaybh = self.date_ngaybaohanh.get_date()

        if not all([mabh, macthd, thoigian]):
            messagebox.showwarning("Cảnh báo", "Vui lòng nhập đầy đủ!")
            return

        try:
            thoigian_int = int(thoigian)
            if thoigian_int <= 0:
                raise ValueError
        except:
            messagebox.showwarning("Cảnh báo", "Thời gian phải là số nguyên dương!")
            return

        old_mabh = self.trHienThi.item(selected[0])["values"][0]
        ngay_het = ngaybh + timedelta(days=thoigian_int * 30)
        trangthai = "CÒN HẠN" if ngay_het >= date.today() else "HẾT HẠN"

        self.trHienThi.item(selected[0], values=(mabh, macthd, self.cb_mahd.get(), thoigian_int, dieukien, self.chuyen_yyyy_sang_dd(ngaybh), trangthai))

        self.ds_sua = [x for x in self.ds_sua if x[0] != old_mabh]
        self.ds_sua.append((mabh, macthd, thoigian_int, dieukien, ngaybh, old_mabh))

        self.xoa_form()
        messagebox.showinfo("Thành công", "Đã sửa! Nhấn 'Lưu' để cập nhật.")

    def xoa(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Cảnh báo", "Chọn dòng cần xóa!")
            return

        if not messagebox.askyesno("Xác nhận", "Xóa dòng này?"):
            return

        mabh = self.trHienThi.item(selected[0])["values"][0]
        self.trHienThi.delete(selected[0])

        self.ds_them = [x for x in self.ds_them if x[0] != mabh]
        if mabh not in self.ds_xoa:
            self.ds_xoa.append(mabh)

        self.xoa_form()
        messagebox.showinfo("Thành công", "Đã xóa! Nhấn 'Lưu' để cập nhật CSDL.")

    def luu(self):
        if not (self.ds_them or self.ds_sua or self.ds_xoa):
            messagebox.showinfo("Thông báo", "Không có thay đổi!")
            return

        if not messagebox.askyesno("Xác nhận", "Lưu tất cả thay đổi?"):
            return

        try:
            for mabh in self.ds_xoa:
                self.cursor.execute("DELETE FROM BaoHanh WHERE MaBH = ?", (mabh,))

            for mabh, macthd, thoigian, dieukien, ngaybh in self.ds_them:
                ngaybh_str = ngaybh.strftime('%Y-%m-%d')
                self.cursor.execute("""
                    INSERT INTO BaoHanh (MaBH, MaCTHD, ThoiGianBaoHanh, DieuKien, NgayBaoHanh)
                    VALUES (?, ?, ?, ?, ?)
                    """,(mabh, macthd, thoigian, dieukien, ngaybh_str),)

            for mabh, macthd, thoigian, dieukien, ngaybh, old_mabh in self.ds_sua:
                ngaybh_str = ngaybh.strftime('%Y-%m-%d')
                self.cursor.execute("""
                    UPDATE BaoHanh SET MaBH=?, MaCTHD=?, ThoiGianBaoHanh=?, DieuKien=?, NgayBaoHanh=?
                    WHERE MaBH=?
                    """,(mabh, macthd, thoigian, dieukien, ngaybh_str, old_mabh),)

            self.conn.commit()
            messagebox.showinfo("Thành công", "Đã lưu tất cả thay đổi!")

        except pyodbc.IntegrityError as e:
            self.conn.rollback()
            messagebox.showerror("Lỗi FK", f"Vi phạm ràng buộc:\n{e}")
        except Exception as e:
            self.conn.rollback()
            messagebox.showerror("Lỗi", f"Lỗi CSDL: {e}")

        self.hienthi_dulieu()
        self.xoa_form()
        self.ds_them.clear()
        self.ds_sua.clear()
        self.ds_xoa.clear()

    def lammoi(self):
        if messagebox.askyesno("Xác nhận", "Hủy tất cả thay đổi?"):
            self.ds_them.clear()
            self.ds_sua.clear()
            self.ds_xoa.clear()
            self.hienthi_dulieu()
            self.xoa_form()
            self.txt_timkiem.delete(0, tk.END)

    def timkiem(self):
        keyword = self.txt_timkiem.get().strip()
        if not keyword:
            self.hienthi_dulieu()
            return

        try:
            for item in self.trHienThi.get_children():
                self.trHienThi.delete(item)

            lua_chon = self.search_option.get()
            sql = ""
            param = f"%{keyword}%"

            if lua_chon == "mabh":
                sql = """
                    SELECT bh.MaBH, bh.MaCTHD, cthd.MaHD, bh.ThoiGianBaoHanh, bh.DieuKien, bh.NgayBaoHanh
                    FROM BaoHanh bh
                    JOIN ChiTietHoaDon cthd ON bh.MaCTHD = cthd.MaCTHD
                    WHERE bh.MaBH LIKE ?
                """
            elif lua_chon == "macthd":
                sql = """
                    SELECT bh.MaBH, bh.MaCTHD, cthd.MaHD, bh.ThoiGianBaoHanh, bh.DieuKien, bh.NgayBaoHanh
                    FROM BaoHanh bh
                    JOIN ChiTietHoaDon cthd ON bh.MaCTHD = cthd.MaCTHD
                    WHERE bh.MaCTHD LIKE ?
                """
            elif lua_chon == "mahd":
                sql = """
                    SELECT bh.MaBH, bh.MaCTHD, cthd.MaHD, bh.ThoiGianBaoHanh, bh.DieuKien, bh.NgayBaoHanh
                    FROM BaoHanh bh
                    JOIN ChiTietHoaDon cthd ON bh.MaCTHD = cthd.MaCTHD
                    WHERE cthd.MaHD LIKE ?
                """

            self.cursor.execute(sql, (param,))
            rows = self.cursor.fetchall()

            for row in rows:
                ngay_bh = row.NgayBaoHanh.date() if hasattr(row.NgayBaoHanh, 'date') else date.fromisoformat(str(row.NgayBaoHanh).split()[0])
                ngay_het = ngay_bh + timedelta(days=row.ThoiGianBaoHanh * 30)
                tt = "CÒN HẠN" if ngay_het >= date.today() else "HẾT HẠN"

                self.trHienThi.insert("", "end", values=(row.MaBH, row.MaCTHD, row.MaHD, row.ThoiGianBaoHanh, row.DieuKien or "", self.chuyen_yyyy_sang_dd(row.NgayBaoHanh), tt))

            if not rows:
                messagebox.showinfo("Thông báo", "Không tìm thấy!")

        except Exception as e:
            messagebox.showerror("Lỗi", f"Tìm kiếm lỗi: {e}")

    def xoa_form(self):
        self.txt_mabh.delete(0, tk.END)
        self.cb_macthd.set("")
        self.cb_mahd.set("")
        self.txt_thoigian.delete(0, tk.END)
        self.txt_dieukien.delete(0, tk.END)
        self.date_ngaybaohanh.set_date(date.today())
        self.lbl_trangthai.config(text="")

    def huy(self):
        self.txt_timkiem.delete(0, tk.END)
        self.hienthi_dulieu()

    def InBaoHanh(self):
        selected = self.trHienThi.selection()
        if not selected:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 bảo hành để in.")
            return

        ma_bh = self.trHienThi.item(selected[0], "values")[0]

        # --- Lấy dữ liệu từ SQL ---
        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT bh.MaBH, bh.MaCTHD, bh.ThoiGianBaoHanh, bh.DieuKien, bh.NgayBaoHanh,
                    cthd.MaHD, tv.MaTivi, tv.TenTivi, kh.MaKH, kh.TenKH
                FROM BaoHanh bh
                JOIN ChiTietHoaDon cthd ON bh.MaCTHD = cthd.MaCTHD
                JOIN HoaDonBan hdb ON cthd.MaHD = hdb.MaHD
                JOIN KhachHang kh ON hdb.MaKH = kh.MaKH
                JOIN Tivi tv ON cthd.MaTivi = tv.MaTivi
                WHERE bh.MaBH = ?
            """, (ma_bh,))
            bh = cursor.fetchone()
            cursor.close()

            if not bh:
                messagebox.showerror("Lỗi", f"Không tìm thấy thông tin bảo hành {ma_bh}")
                return

        except Exception as e:
            messagebox.showerror("Lỗi CSDL", f"Không thể lấy dữ liệu để in phiếu bảo hành:\n{e}")
            return

        # --- Tạo file Word ---
        try:
            document = Document()

            # Cấu hình font mặc định
            style = document.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(11)

            # --- Tiêu đề ---
            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("PHIẾU BẢO HÀNH SẢN PHẨM")
            run.bold = True
            run.font.size = Pt(16)
            document.add_paragraph()

            # --- Thông tin chung ---
            TAB_POS = Cm(8)

            p1 = document.add_paragraph()
            p1.paragraph_format.tab_stops.add_tab_stop(TAB_POS, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p1.add_run("Mã bảo hành: ").bold = True
            run = p1.add_run(bh.MaBH)
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
            p1.add_run("\t")
            p1.add_run("Mã chi tiết hóa đơn: ").bold = True
            p1.add_run(bh.MaCTHD)

            p2 = document.add_paragraph()
            p2.paragraph_format.tab_stops.add_tab_stop(TAB_POS, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p2.add_run("Mã hóa đơn: ").bold = True
            p2.add_run(bh.MaHD)
            p2.add_run("\t")
            p2.add_run("Mã khách hàng: ").bold = True
            p2.add_run(bh.MaKH)

            p3 = document.add_paragraph()
            p3.paragraph_format.tab_stops.add_tab_stop(TAB_POS, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p3.add_run("Tên khách hàng: ").bold = True
            p3.add_run(bh.TenKH)
            p3.add_run("\t")
            p3.add_run("Sản phẩm: ").bold = True
            p3.add_run(f"{bh.MaTivi} - {bh.TenTivi}")

            document.add_paragraph()

            # --- Thông tin bảo hành ---
            ngay_bh = bh.NgayBaoHanh.date() if hasattr(bh.NgayBaoHanh, 'date') else datetime.strptime(str(bh.NgayBaoHanh).split()[0], "%Y-%m-%d").date()
            ngay_het = ngay_bh + timedelta(days=bh.ThoiGianBaoHanh * 30)
            tt = "CÒN HẠN" if ngay_het >= datetime.today().date() else "HẾT HẠN"

            document.add_paragraph(f"Ngày bảo hành: {ngay_bh.strftime('%d/%m/%Y')}")
            document.add_paragraph(f"Thời gian bảo hành: {bh.ThoiGianBaoHanh} tháng")
            document.add_paragraph(f"Ngày hết hạn: {ngay_het.strftime('%d/%m/%Y')}")
            document.add_paragraph(f"Tình trạng: {tt}")
            document.add_paragraph(f"Điều kiện bảo hành: {bh.DieuKien or 'Không có'}")

            document.add_paragraph()
            document.add_paragraph("Sản phẩm sẽ được bảo hành miễn phí nếu đáp ứng các điều kiện nêu trên.", style="Normal")

            # --- Khu vực ký tên ---
            document.add_paragraph()
            document.add_paragraph()
            sig_table = document.add_table(rows=2, cols=3)
            sig_table.style = None
            sig_table.autofit = False

            headers = ["Người lập phiếu", "Kỹ thuật viên", "Khách hàng"]
            for i, h in enumerate(headers):
                cell = sig_table.cell(0, i)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h)
                run.bold = True

            for i in range(3):
                cell = sig_table.cell(1, i)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run("(Ký, ghi rõ họ tên)")

            # --- Chọn nơi lưu file ---
            default_name = f"PhieuBaoHanh_{ma_bh}.docx"
            file_path = filedialog.asksaveasfilename(
                title="Chọn nơi lưu phiếu bảo hành",
                defaultextension=".docx",
                initialfile=default_name,
                filetypes=[("Word Document", "*.docx")]
            )

            if not file_path:
                messagebox.showinfo("Thông báo", "Đã hủy lưu file.")
                return

            # --- Lưu file ---
            document.save(file_path)
            messagebox.showinfo("Thành công", f"Đã tạo phiếu bảo hành tại:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi khi tạo file Word:\n{e}")
